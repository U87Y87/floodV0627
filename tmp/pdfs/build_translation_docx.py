import re
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"D:\Debris Flow\V0622")
ZH = ROOT / "tmp/pdfs/paper_text/full_zh.txt"
EN_PAGES = ROOT / "tmp/pdfs/paper_text"
OUT = ROOT / "output/documents/多源时空数据立方体_中文译稿.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)


def set_font(run, east="宋体", latin="Times New Roman", size=11, bold=None, color=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def clean_zh(text):
    replacements = {
        "大地球观测": "大规模对地观测",
        "地球观测": "对地观测",
        "光电数据": "对地观测数据",
        "光电领域": "对地观测领域",
        "范例": "范式",
        "网络基础设施": "网络信息基础设施",
        "图块": "瓦片",
        "用于大规模地理空间分析的多源时空数据立方体": "面向大规模地理空间分析的多源时空数据立方体",
        "高范": "高帆",
        "赵帅锋": "赵帅峰",
        "姜村": "蒋良村",
        "上官博一": "上官博一",
        "数据立方体的容量": "数据立方体的能力",
        "国际地理信息科学杂志": "《国际地理信息科学杂志》",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\s+([，。；：！？）])", r"\1", text)
    text = re.sub(r"（\s+", "（", text)
    return text.strip()


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = Inches(1)
sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333
for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK, 8, 4),
]:
    st = styles[name]
    st.font.name = "Arial"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = header.add_run("学术文献中文译稿 | GeoCube")
set_font(r, east="微软雅黑", latin="Arial", size=8.5, color=GRAY)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("第 ")
set_font(r, size=9, color=GRAY)
add_page_field(footer)
r = footer.add_run(" 页")
set_font(r, size=9, color=GRAY)

# Editorial cover.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(95)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("研究论文 · 中文译稿")
set_font(r, east="微软雅黑", latin="Arial", size=11, bold=True, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
r = p.add_run("面向大规模地理空间分析的\n多源时空数据立方体")
set_font(r, east="微软雅黑", latin="Arial", size=24, bold=True, color=DARK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
r = p.add_run("A multi-source spatio-temporal data cube for large-scale geospatial analysis")
set_font(r, east="微软雅黑", latin="Arial", size=12, color=GRAY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("高帆、岳鹏、曹志鹏、赵帅峰、上官博一、蒋良村、胡磊、方哲、梁哲恒")
set_font(r, east="宋体", latin="Times New Roman", size=11, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("International Journal of Geographical Information Science, 36(9), 1853-1884 (2022)\nDOI: 10.1080/13658816.2022.2087222")
set_font(r, size=10, color=GRAY)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(55)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("说明：本稿为辅助阅读的中文译文。公式、缩写、数据与参考文献请以英文原文为准。")
set_font(r, east="楷体", latin="Times New Roman", size=9.5, color=GRAY)
doc.add_page_break()

raw = ZH.read_text(encoding="utf-8")
# Start with article page 2 and keep translated body through contributor material.
raw = raw.split("===== PAGE 2 =====", 1)[1]
body, _translated_refs = raw.split("===== PAGE 31 =====", 1)
blocks = [clean_zh(x) for x in re.split(r"\n\s*\n", body) if x.strip()]

for block in blocks:
    if block.startswith("===== PAGE"):
        continue
    if re.fullmatch(r"[一二三四五六七八九十]+、.+", block) or re.match(r"^\d+\.\s*[^。]{1,60}$", block):
        p = doc.add_paragraph(style="Heading 1")
    elif re.match(r"^\d+\.\d+\.\s*[^。]{1,80}$", block):
        p = doc.add_paragraph(style="Heading 2")
    elif block in {"摘要", "关键词", "致谢", "利益披露声明", "作者简介", "ORCID", "数据和代码可用性声明"}:
        p = doc.add_paragraph(style="Heading 1")
    elif re.match(r"^(图|表)\s*\d+", block):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.29)
    r = p.add_run(block)
    set_font(r, size=11)
    if re.match(r"^(图|表)\s*\d+", block):
        r.bold = True

doc.add_paragraph("参考文献", style="Heading 1")
refs = []
page30 = (EN_PAGES / "page-30.txt").read_text(encoding="utf-8")
refs.append(page30.split("References", 1)[1])
refs.extend((EN_PAGES / f"page-{n:02d}.txt").read_text(encoding="utf-8") for n in (31, 32, 33))
ref_text = "\n".join(refs)
ref_text = re.sub(r"\n(?=[a-z\[\(])", " ", ref_text)
ref_text = re.sub(r"\n+", "\n", ref_text)
ref_text = re.sub(r"INTERNATIONAL JOURNAL.*?\d{4}\s*", "", ref_text)
ref_text = re.sub(r"\n\d{4}\s+F\.\s+GAO.*?\n", "\n", ref_text)
for line in ref_text.splitlines():
    line = line.strip()
    if not line or re.fullmatch(r"\d{4}", line):
        continue
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(line)
    set_font(r, size=9)

doc.core_properties.title = "面向大规模地理空间分析的多源时空数据立方体 - 中文译稿"
doc.core_properties.subject = "学术论文中文翻译"
doc.core_properties.author = "中文译稿"
doc.save(OUT)
print(OUT)
